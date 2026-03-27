#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Модуль для создания документов Word (.docx)
"""

import logging
from io import BytesIO
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Класс для создания документов Word"""
    
    def __init__(self):
        pass
    
    def create_document(self, title: str, content: str, style: str = "default") -> BytesIO:
        """
        Создать документ Word
        
        Args:
            title: Заголовок документа
            content: Содержимое документа
            style: Стиль документа (default, formal, summary)
            
        Returns:
            BytesIO объект с документом
        """
        try:
            doc = Document()
            
            # Настройка стилей
            if style == "formal":
                self._apply_formal_style(doc, title, content)
            elif style == "summary":
                self._apply_summary_style(doc, title, content)
            else:
                self._apply_default_style(doc, title, content)
            
            # Сохраняем в BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            logger.info(f"Создан документ Word: {title}, стиль: {style}")
            return buffer
            
        except Exception as e:
            logger.error(f"Ошибка создания документа Word: {e}")
            raise
    
    def _apply_default_style(self, doc: Document, title: str, content: str):
        """Применить стиль по умолчанию"""
        # Заголовок
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Содержимое
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)
    
    def _apply_formal_style(self, doc: Document, title: str, content: str):
        """Применить формальный стиль"""
        # Заголовок
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата (если есть в контенте)
        from datetime import datetime
        date_para = doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.paragraph_format.space_after = Pt(12)
        
        # Содержимое
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(6)
    
    def _apply_summary_style(self, doc: Document, title: str, content: str):
        """Применить стиль конспекта"""
        # Заголовок
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Подзаголовок
        subtitle = doc.add_paragraph("Конспект")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.italic = True
        subtitle_format.size = Pt(12)
        subtitle.paragraph_format.space_after = Pt(12)
        
        # Содержимое с форматированием
        paragraphs = content.split('\n')
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Если строка начинается с цифры или маркера - это заголовок раздела
            if para.startswith(('1.', '2.', '3.', '4.', '5.', '•', '-', '*')):
                p = doc.add_heading(para, level=2)
            elif para.startswith(('**', '__')) and para.endswith(('**', '__')):
                # Жирный текст
                p = doc.add_paragraph()
                run = p.add_run(para.strip('*_'))
                run.bold = True
            else:
                p = doc.add_paragraph(para)
            
            p.paragraph_format.space_after = Pt(6)
    
    def create_summary(self, text: str, summary_type: str = "краткий") -> str:
        """
        Создать конспект из текста (заглушка, реальная обработка через Claude)
        
        Args:
            text: Исходный текст
            summary_type: Тип конспекта (краткий, подробный, структурированный)
            
        Returns:
            Текст конспекта
        """
        # Это заглушка, реальная обработка будет через Claude
        if summary_type == "краткий":
            return f"Краткий конспект:\n\n{text[:500]}..."
        elif summary_type == "подробный":
            return f"Подробный конспект:\n\n{text}"
        else:
            return f"Структурированный конспект:\n\n{text}"

