#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Конвертер файлов для Telegram бота
Поддерживает конвертацию между различными форматами документов
"""

import os
import io
import logging
from typing import Dict, List, Optional, Tuple
import tempfile

# PDF обработка
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word обработка
try:
    from docx import Document
    from docx.shared import Inches
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# Excel обработка
try:
    import pandas as pd
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Другие форматы
try:
    import docx2txt
    import markdown
    import html2text
    OTHER_FORMATS_AVAILABLE = True
except ImportError:
    OTHER_FORMATS_AVAILABLE = False

logger = logging.getLogger(__name__)

class FileConverter:
    """Класс для конвертации файлов между различными форматами"""
    
    def __init__(self):
        self.supported_formats = {
            # Документы
            'pdf': ['docx', 'txt', 'html', 'md', 'rtf', 'xlsx'],
            'docx': ['pdf', 'txt', 'html', 'md', 'rtf', 'odt'],
            'doc': ['docx', 'txt', 'html', 'md', 'rtf'],
            'odt': ['docx', 'txt', 'html', 'md', 'rtf'],
            'rtf': ['docx', 'txt', 'html', 'md'],
            
            # Текстовые форматы
            'txt': ['pdf', 'docx', 'html', 'md', 'rtf'],
            'html': ['pdf', 'docx', 'txt', 'md', 'rtf'],
            'htm': ['pdf', 'docx', 'txt', 'md', 'rtf'],
            'md': ['pdf', 'docx', 'txt', 'html', 'rtf'],
            'markdown': ['pdf', 'docx', 'txt', 'html', 'rtf'],
            
            # Таблицы
            'xlsx': ['csv', 'json', 'html', 'pdf', 'ods'],
            'xls': ['xlsx', 'csv', 'json', 'html', 'pdf'],
            'csv': ['xlsx', 'json', 'html', 'pdf', 'ods'],
            'ods': ['xlsx', 'csv', 'json', 'html'],
            
            # Данные
            'json': ['xlsx', 'csv', 'html', 'xml', 'yaml'],
            'xml': ['json', 'html', 'txt'],
            'yaml': ['json', 'xml', 'txt'],
            'yml': ['json', 'xml', 'txt'],
            
            # Презентации
            'pptx': ['pdf', 'html', 'txt'],
            'ppt': ['pptx', 'pdf', 'html', 'txt'],
            
            # Изображения
            'png': ['jpg', 'jpeg', 'webp', 'bmp', 'tiff'],
            'jpg': ['png', 'webp', 'bmp', 'tiff'],
            'jpeg': ['png', 'webp', 'bmp', 'tiff'],
            'webp': ['png', 'jpg', 'jpeg', 'bmp'],
            'bmp': ['png', 'jpg', 'jpeg', 'webp'],
            'tiff': ['png', 'jpg', 'jpeg', 'webp'],
            'gif': ['png', 'jpg', 'jpeg', 'webp'],
            
            # Аудио
            'mp3': ['wav', 'ogg', 'flac', 'aac'],
            'wav': ['mp3', 'ogg', 'flac', 'aac'],
            'ogg': ['mp3', 'wav', 'flac', 'aac'],
            'flac': ['mp3', 'wav', 'ogg', 'aac'],
            'aac': ['mp3', 'wav', 'ogg', 'flac'],
            
            # Видео
            'mp4': ['avi', 'mov', 'mkv', 'webm'],
            'avi': ['mp4', 'mov', 'mkv', 'webm'],
            'mov': ['mp4', 'avi', 'mkv', 'webm'],
            'mkv': ['mp4', 'avi', 'mov', 'webm'],
            'webm': ['mp4', 'avi', 'mov', 'mkv'],
            
            # Архивы
            'zip': ['rar', '7z', 'tar', 'gz'],
            'rar': ['zip', '7z', 'tar', 'gz'],
            '7z': ['zip', 'rar', 'tar', 'gz'],
            'tar': ['zip', 'rar', '7z', 'gz'],
            'gz': ['zip', 'rar', '7z', 'tar'],
            
            # Код
            'py': ['txt', 'html', 'md'],
            'js': ['txt', 'html', 'md'],
            'ts': ['txt', 'html', 'md'],
            'java': ['txt', 'html', 'md'],
            'cpp': ['txt', 'html', 'md'],
            'c': ['txt', 'html', 'md'],
            'cs': ['txt', 'html', 'md'],
            'php': ['txt', 'html', 'md'],
            'rb': ['txt', 'html', 'md'],
            'go': ['txt', 'html', 'md'],
            'rs': ['txt', 'html', 'md'],
            'swift': ['txt', 'html', 'md'],
            'kt': ['txt', 'html', 'md'],
            'sh': ['txt', 'html', 'md'],
            'ps1': ['txt', 'html', 'md'],
            'sql': ['txt', 'html', 'md'],
            'css': ['txt', 'html', 'md'],
            'scss': ['txt', 'html', 'md'],
            'sass': ['txt', 'html', 'md'],
            'less': ['txt', 'html', 'md']
        }
        
        self.format_descriptions = {
            # Документы
            'pdf': 'PDF документ',
            'docx': 'Word документ',
            'doc': 'Word документ (старый)',
            'odt': 'OpenDocument текст',
            'rtf': 'Rich Text Format',
            
            # Текстовые форматы
            'txt': 'Текстовый файл',
            'html': 'HTML страница',
            'htm': 'HTML страница',
            'md': 'Markdown документ',
            'markdown': 'Markdown документ',
            
            # Таблицы
            'xlsx': 'Excel таблица',
            'xls': 'Excel таблица (старый)',
            'csv': 'CSV таблица',
            'ods': 'OpenDocument таблица',
            
            # Данные
            'json': 'JSON данные',
            'xml': 'XML документ',
            'yaml': 'YAML конфигурация',
            'yml': 'YAML конфигурация',
            
            # Презентации
            'pptx': 'PowerPoint презентация',
            'ppt': 'PowerPoint презентация (старая)',
            
            # Изображения
            'png': 'PNG изображение',
            'jpg': 'JPEG изображение',
            'jpeg': 'JPEG изображение',
            'webp': 'WebP изображение',
            'bmp': 'BMP изображение',
            'tiff': 'TIFF изображение',
            'gif': 'GIF анимация',
            
            # Аудио
            'mp3': 'MP3 аудио',
            'wav': 'WAV аудио',
            'ogg': 'OGG аудио',
            'flac': 'FLAC аудио',
            'aac': 'AAC аудио',
            
            # Видео
            'mp4': 'MP4 видео',
            'avi': 'AVI видео',
            'mov': 'MOV видео',
            'mkv': 'MKV видео',
            'webm': 'WebM видео',
            
            # Архивы
            'zip': 'ZIP архив',
            'rar': 'RAR архив',
            '7z': '7-Zip архив',
            'tar': 'TAR архив',
            'gz': 'GZIP архив',
            
            # Код
            'py': 'Python код',
            'js': 'JavaScript код',
            'ts': 'TypeScript код',
            'java': 'Java код',
            'cpp': 'C++ код',
            'c': 'C код',
            'cs': 'C# код',
            'php': 'PHP код',
            'rb': 'Ruby код',
            'go': 'Go код',
            'rs': 'Rust код',
            'swift': 'Swift код',
            'kt': 'Kotlin код',
            'sh': 'Bash скрипт',
            'ps1': 'PowerShell скрипт',
            'sql': 'SQL запрос',
            'css': 'CSS стили',
            'scss': 'SCSS стили',
            'sass': 'SASS стили',
            'less': 'LESS стили'
        }
    
    def detect_file_type(self, filename: str) -> str:
        """Определение типа файла по расширению"""
        extension = filename.lower().split('.')[-1] if '.' in filename else ''
        
        # Нормализация расширений
        extension_map = {
            'doc': 'docx',
            'docm': 'docx',
            'xls': 'xlsx',
            'xlsm': 'xlsx'
        }
        
        return extension_map.get(extension, extension)
    
    def is_supported_format(self, file_format: str) -> bool:
        """Проверка поддержки формата"""
        return file_format.lower() in self.supported_formats
    
    def get_available_conversions(self, source_format: str) -> List[str]:
        """Получение доступных форматов для конвертации"""
        return self.supported_formats.get(source_format.lower(), [])
    
    def convert_pdf_to_text(self, pdf_content: bytes) -> str:
        """Конвертация PDF в текст"""
        if not PDF_AVAILABLE:
            raise Exception("PDF обработка недоступна. Установите PyPDF2 и pdfplumber")
        
        text_content = ""
        
        try:
            # Пробуем pdfplumber (лучше для сложных PDF)
            with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content += text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
            
            # Fallback на PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content += text + "\n"
            except Exception as e2:
                raise Exception(f"Не удалось извлечь текст из PDF: {e2}")
        
        return text_content.strip()
    
    def convert_pdf_to_excel(self, pdf_content: bytes, filename: str = "converted.xlsx") -> bytes:
        """Конвертация PDF в Excel с извлечением таблиц"""
        if not PDF_AVAILABLE or not EXCEL_AVAILABLE:
            raise Exception("Для конвертации PDF в Excel требуется PyPDF2, pdfplumber, pandas и openpyxl")
        
        try:
            import pandas as pd
            
            # Извлекаем текст и таблицы из PDF
            text_content = ""
            tables_data = []
            
            with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Извлекаем текст
                    text = page.extract_text()
                    if text:
                        text_content += f"Страница {page_num + 1}:\n{text}\n\n"
                    
                    # Извлекаем таблицы
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables):
                        if table:
                            # Конвертируем таблицу в DataFrame
                            df = pd.DataFrame(table[1:], columns=table[0] if table[0] else None)
                            tables_data.append({
                                'page': page_num + 1,
                                'table': table_num + 1,
                                'data': df
                            })
            
            # Создаем Excel файл
            excel_buffer = io.BytesIO()
            
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                # Добавляем лист с текстом
                if text_content:
                    text_df = pd.DataFrame([text_content], columns=['Текст из PDF'])
                    text_df.to_excel(writer, sheet_name='Текст', index=False)
                
                # Добавляем листы с таблицами
                for i, table_info in enumerate(tables_data):
                    sheet_name = f"Таблица_{table_info['page']}_{table_info['table']}"
                    table_info['data'].to_excel(writer, sheet_name=sheet_name, index=False)
                
                # Если нет таблиц, создаем лист с общим текстом
                if not tables_data and text_content:
                    # Разбиваем текст на строки для лучшего отображения
                    lines = text_content.split('\n')
                    text_df = pd.DataFrame(lines, columns=['Строка'])
                    text_df.to_excel(writer, sheet_name='Содержимое', index=False)
            
            excel_buffer.seek(0)
            return excel_buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"Ошибка конвертации PDF в Excel: {e}")
    
    def convert_pdf_to_word(self, pdf_content: bytes, filename: str = "converted.docx") -> bytes:
        """Конвертация PDF в Word документ"""
        if not PDF_AVAILABLE or not WORD_AVAILABLE:
            raise Exception("Для конвертации PDF в Word требуется PyPDF2, pdfplumber и python-docx")
        
        try:
            # Извлекаем текст из PDF
            text_content = self.convert_pdf_to_text(pdf_content)
            
            # Создаем Word документ
            doc = Document()
            
            # Добавляем заголовок
            doc.add_heading('Конвертированный PDF документ', 0)
            
            # Разбиваем текст на параграфы
            paragraphs = text_content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Если строка короткая и содержит только заглавные буквы - делаем заголовком
                    if len(para_text.strip()) < 100 and para_text.strip().isupper():
                        doc.add_heading(para_text.strip(), level=1)
                    else:
                        # Разбиваем длинные параграфы на предложения
                        sentences = para_text.split('. ')
                        for sentence in sentences:
                            if sentence.strip():
                                doc.add_paragraph(sentence.strip() + ('.' if not sentence.endswith('.') else ''))
            
            # Сохраняем в байты
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"Ошибка конвертации PDF в Word: {e}")
    
    def convert_text_to_docx(self, text_content: str, filename: str = "converted.docx") -> bytes:
        """Конвертация текста в Word документ"""
        if not WORD_AVAILABLE:
            raise Exception("Word обработка недоступна. Установите python-docx")
        
        doc = Document()
        
        # Добавляем заголовок
        doc.add_heading(f'Конвертированный документ', 0)
        
        # Разбиваем текст на параграфы
        paragraphs = text_content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                # Если строка короткая и содержит только заглавные буквы - делаем заголовком
                if len(para_text.strip()) < 100 and para_text.strip().isupper():
                    doc.add_heading(para_text.strip(), level=1)
                else:
                    doc.add_paragraph(para_text.strip())
        
        # Сохраняем в байты
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
    
    def convert_text_to_html(self, text_content: str) -> str:
        """Конвертация текста в HTML"""
        # Простая конвертация в HTML
        html_content = f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Конвертированный документ</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1 {{ color: #333; border-bottom: 2px solid #333; }}
        p {{ margin-bottom: 15px; }}
    </style>
</head>
<body>
    <h1>Конвертированный документ</h1>
"""
        
        # Разбиваем на параграфы
        paragraphs = text_content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                if len(para_text.strip()) < 100 and para_text.strip().isupper():
                    html_content += f"    <h2>{para_text.strip()}</h2>\n"
                else:
                    html_content += f"    <p>{para_text.strip()}</p>\n"
        
        html_content += "</body>\n</html>"
        
        return html_content
    
    def convert_text_to_markdown(self, text_content: str) -> str:
        """Конвертация текста в Markdown"""
        md_content = "# Конвертированный документ\n\n"
        
        # Разбиваем на параграфы
        paragraphs = text_content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                if len(para_text.strip()) < 100 and para_text.strip().isupper():
                    md_content += f"## {para_text.strip()}\n\n"
                else:
                    md_content += f"{para_text.strip()}\n\n"
        
        return md_content
    
    def convert_html_to_text(self, html_content: str) -> str:
        """Конвертация HTML в текст"""
        if OTHER_FORMATS_AVAILABLE:
            try:
                h = html2text.HTML2Text()
                h.ignore_links = False
                return h.handle(html_content)
            except:
                pass
        
        # Простая конвертация без библиотек
        import re
        
        # Удаляем HTML теги
        text = re.sub(r'<[^>]+>', '', html_content)
        
        # Декодируем HTML entities
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&amp;', '&')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&quot;', '"')
        
        return text.strip()
    
    def convert_markdown_to_text(self, md_content: str) -> str:
        """Конвертация Markdown в текст"""
        if OTHER_FORMATS_AVAILABLE:
            try:
                html = markdown.markdown(md_content)
                return self.convert_html_to_text(html)
            except:
                pass
        
        # Простая конвертация без библиотек
        import re
        
        # Удаляем markdown разметку
        text = re.sub(r'#{1,6}\s*', '', md_content)  # Заголовки
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Жирный текст
        text = re.sub(r'\*(.*?)\*', r'\1', text)      # Курсив
        text = re.sub(r'`(.*?)`', r'\1', text)        # Код
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Ссылки
        
        return text.strip()
    
    def convert_image_format(self, image_content: bytes, source_format: str, target_format: str) -> bytes:
        """Конвертация изображений между форматами"""
        try:
            from PIL import Image
            
            # Открываем изображение
            image = Image.open(io.BytesIO(image_content))
            
            # Конвертируем в RGB если нужно (для JPEG)
            if target_format.lower() in ['jpg', 'jpeg'] and image.mode in ['RGBA', 'LA', 'P']:
                # Создаем белый фон для прозрачных изображений
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                image = background
            elif target_format.lower() in ['png', 'webp'] and image.mode == 'RGB':
                image = image.convert('RGBA')
            
            # Сохраняем в новом формате
            output_buffer = io.BytesIO()
            
            if target_format.lower() == 'jpg' or target_format.lower() == 'jpeg':
                image.save(output_buffer, format='JPEG', quality=95)
            elif target_format.lower() == 'png':
                image.save(output_buffer, format='PNG')
            elif target_format.lower() == 'webp':
                image.save(output_buffer, format='WEBP', quality=95)
            elif target_format.lower() == 'bmp':
                image.save(output_buffer, format='BMP')
            elif target_format.lower() == 'tiff':
                image.save(output_buffer, format='TIFF')
            else:
                raise Exception(f"Неподдерживаемый формат изображения: {target_format}")
            
            output_buffer.seek(0)
            return output_buffer.getvalue()
            
        except ImportError:
            raise Exception("Для конвертации изображений требуется библиотека Pillow")
        except Exception as e:
            raise Exception(f"Ошибка конвертации изображения: {e}")
    
    def convert_code_to_html(self, code_content: str, language: str) -> str:
        """Конвертация кода в HTML с подсветкой синтаксиса"""
        try:
            import pygments
            from pygments import highlight
            from pygments.lexers import get_lexer_by_name
            from pygments.formatters import HtmlFormatter
            
            # Получаем лексер для языка
            lexer = get_lexer_by_name(language, stripall=True)
            
            # Создаем HTML форматтер
            formatter = HtmlFormatter(
                style='default',
                linenos=True,
                cssclass='highlight',
                wrapcode=True
            )
            
            # Генерируем HTML
            html_code = highlight(code_content, lexer, formatter)
            
            # Добавляем CSS стили
            css_styles = formatter.get_style_defs('.highlight')
            
            full_html = f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Код - {language}</title>
    <style>
        body {{ font-family: 'Courier New', monospace; margin: 20px; background: #f5f5f5; }}
        .container {{ background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        h1 {{ color: #333; margin-bottom: 20px; }}
        {css_styles}
    </style>
</head>
<body>
    <div class="container">
        <h1>Код ({language})</h1>
        {html_code}
    </div>
</body>
</html>"""
            
            return full_html
            
        except ImportError:
            # Fallback без подсветки синтаксиса
            escaped_code = code_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <title>Код - {language}</title>
    <style>
        body {{ font-family: 'Courier New', monospace; margin: 20px; background: #f5f5f5; }}
        .container {{ background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        h1 {{ color: #333; }}
        pre {{ background: #f8f8f8; padding: 15px; border-radius: 4px; overflow-x: auto; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Код ({language})</h1>
        <pre><code>{escaped_code}</code></pre>
    </div>
</body>
</html>"""
        except Exception as e:
            raise Exception(f"Ошибка конвертации кода в HTML: {e}")
    
    def convert_excel_to_csv(self, excel_content: bytes) -> str:
        """Конвертация Excel в CSV"""
        if not EXCEL_AVAILABLE:
            raise Exception("Excel обработка недоступна. Установите pandas и openpyxl")
        
        try:
            # Читаем Excel файл
            df = pd.read_excel(io.BytesIO(excel_content))
            
            # Конвертируем в CSV
            csv_buffer = io.StringIO()
            df.to_csv(csv_buffer, index=False, encoding='utf-8')
            
            return csv_buffer.getvalue()
        except Exception as e:
            raise Exception(f"Ошибка конвертации Excel в CSV: {e}")
    
    def convert_csv_to_excel(self, csv_content: str) -> bytes:
        """Конвертация CSV в Excel"""
        if not EXCEL_AVAILABLE:
            raise Exception("Excel обработка недоступна. Установите pandas и openpyxl")
        
        try:
            # Читаем CSV
            df = pd.read_csv(io.StringIO(csv_content))
            
            # Конвертируем в Excel
            excel_buffer = io.BytesIO()
            df.to_excel(excel_buffer, index=False, engine='openpyxl')
            excel_buffer.seek(0)
            
            return excel_buffer.getvalue()
        except Exception as e:
            raise Exception(f"Ошибка конвертации CSV в Excel: {e}")
    
    def convert_json_to_excel(self, json_content: str) -> bytes:
        """Конвертация JSON в Excel"""
        if not EXCEL_AVAILABLE:
            raise Exception("Excel обработка недоступна. Установите pandas и openpyxl")
        
        try:
            import json
            
            # Парсим JSON
            data = json.loads(json_content)
            
            # Конвертируем в DataFrame
            if isinstance(data, list):
                df = pd.DataFrame(data)
            elif isinstance(data, dict):
                df = pd.DataFrame([data])
            else:
                raise Exception("Неподдерживаемый формат JSON")
            
            # Конвертируем в Excel
            excel_buffer = io.BytesIO()
            df.to_excel(excel_buffer, index=False, engine='openpyxl')
            excel_buffer.seek(0)
            
            return excel_buffer.getvalue()
        except Exception as e:
            raise Exception(f"Ошибка конвертации JSON в Excel: {e}")
    
    def convert_file(self, file_content: bytes, source_format: str, target_format: str, filename: str = "") -> Tuple[bytes, str]:
        """
        Основная функция конвертации файлов
        
        Args:
            file_content: Содержимое исходного файла
            source_format: Исходный формат
            target_format: Целевой формат
            filename: Имя файла (опционально)
        
        Returns:
            Tuple[bytes, str]: (конвертированное содержимое, новое имя файла)
        """
        source_format = source_format.lower()
        target_format = target_format.lower()
        
        if not self.is_supported_format(source_format):
            raise Exception(f"Неподдерживаемый исходный формат: {source_format}")
        
        if target_format not in self.get_available_conversions(source_format):
            raise Exception(f"Конвертация из {source_format} в {target_format} не поддерживается")
        
        # Определяем новое имя файла
        if not filename:
            filename = f"converted.{target_format}"
        else:
            name, _ = os.path.splitext(filename)
            filename = f"{name}.{target_format}"
        
        # Выполняем конвертацию
        if source_format == 'pdf':
            if target_format == 'txt':
                text_content = self.convert_pdf_to_text(file_content)
                return text_content.encode('utf-8'), filename
            elif target_format == 'docx':
                docx_content = self.convert_pdf_to_word(file_content, filename)
                return docx_content, filename
            elif target_format == 'xlsx':
                excel_content = self.convert_pdf_to_excel(file_content, filename)
                return excel_content, filename
            elif target_format == 'html':
                text_content = self.convert_pdf_to_text(file_content)
                html_content = self.convert_text_to_html(text_content)
                return html_content.encode('utf-8'), filename
            elif target_format == 'md':
                text_content = self.convert_pdf_to_text(file_content)
                md_content = self.convert_text_to_markdown(text_content)
                return md_content.encode('utf-8'), filename
        
        elif source_format == 'docx':
            if target_format == 'txt':
                # Для docx нужна специальная обработка
                raise Exception("Конвертация DOCX в TXT пока не поддерживается")
            elif target_format == 'pdf':
                raise Exception("Конвертация DOCX в PDF пока не поддерживается")
            elif target_format == 'html':
                raise Exception("Конвертация DOCX в HTML пока не поддерживается")
            elif target_format == 'md':
                raise Exception("Конвертация DOCX в Markdown пока не поддерживается")
        
        elif source_format == 'txt':
            text_content = file_content.decode('utf-8')
            if target_format == 'docx':
                docx_content = self.convert_text_to_docx(text_content, filename)
                return docx_content, filename
            elif target_format == 'html':
                html_content = self.convert_text_to_html(text_content)
                return html_content.encode('utf-8'), filename
            elif target_format == 'md':
                md_content = self.convert_text_to_markdown(text_content)
                return md_content.encode('utf-8'), filename
        
        elif source_format == 'html':
            html_content = file_content.decode('utf-8')
            if target_format == 'txt':
                text_content = self.convert_html_to_text(html_content)
                return text_content.encode('utf-8'), filename
            elif target_format == 'docx':
                text_content = self.convert_html_to_text(html_content)
                docx_content = self.convert_text_to_docx(text_content, filename)
                return docx_content, filename
            elif target_format == 'md':
                text_content = self.convert_html_to_text(html_content)
                md_content = self.convert_text_to_markdown(text_content)
                return md_content.encode('utf-8'), filename
        
        elif source_format == 'md':
            md_content = file_content.decode('utf-8')
            if target_format == 'txt':
                text_content = self.convert_markdown_to_text(md_content)
                return text_content.encode('utf-8'), filename
            elif target_format == 'docx':
                text_content = self.convert_markdown_to_text(md_content)
                docx_content = self.convert_text_to_docx(text_content, filename)
                return docx_content, filename
            elif target_format == 'html':
                text_content = self.convert_markdown_to_text(md_content)
                html_content = self.convert_text_to_html(text_content)
                return html_content.encode('utf-8'), filename
        
        elif source_format == 'xlsx':
            if target_format == 'csv':
                csv_content = self.convert_excel_to_csv(file_content)
                return csv_content.encode('utf-8'), filename
            elif target_format == 'json':
                # Конвертация Excel в JSON
                df = pd.read_excel(io.BytesIO(file_content))
                json_content = df.to_json(orient='records', force_ascii=False, indent=2)
                return json_content.encode('utf-8'), filename
            elif target_format == 'html':
                df = pd.read_excel(io.BytesIO(file_content))
                html_content = df.to_html(index=False)
                return html_content.encode('utf-8'), filename
        
        elif source_format == 'csv':
            csv_content = file_content.decode('utf-8')
            if target_format == 'xlsx':
                excel_content = self.convert_csv_to_excel(csv_content)
                return excel_content, filename
            elif target_format == 'json':
                df = pd.read_csv(io.StringIO(csv_content))
                json_content = df.to_json(orient='records', force_ascii=False, indent=2)
                return json_content.encode('utf-8'), filename
            elif target_format == 'html':
                df = pd.read_csv(io.StringIO(csv_content))
                html_content = df.to_html(index=False)
                return html_content.encode('utf-8'), filename
        
        elif source_format == 'json':
            json_content = file_content.decode('utf-8')
            if target_format == 'xlsx':
                excel_content = self.convert_json_to_excel(json_content)
                return excel_content, filename
            elif target_format == 'csv':
                data = json.loads(json_content)
                if isinstance(data, list):
                    df = pd.DataFrame(data)
                elif isinstance(data, dict):
                    df = pd.DataFrame([data])
                else:
                    raise Exception("Неподдерживаемый формат JSON")
                csv_content = df.to_csv(index=False)
                return csv_content.encode('utf-8'), filename
            elif target_format == 'html':
                data = json.loads(json_content)
                if isinstance(data, list):
                    df = pd.DataFrame(data)
                elif isinstance(data, dict):
                    df = pd.DataFrame([data])
                else:
                    raise Exception("Неподдерживаемый формат JSON")
                html_content = df.to_html(index=False)
                return html_content.encode('utf-8'), filename
        
        # Обработка изображений
        elif source_format in ['png', 'jpg', 'jpeg', 'webp', 'bmp', 'tiff', 'gif']:
            if target_format in ['png', 'jpg', 'jpeg', 'webp', 'bmp', 'tiff']:
                converted_content = self.convert_image_format(file_content, source_format, target_format)
                return converted_content, filename
        
        # Обработка кода
        elif source_format in ['py', 'js', 'ts', 'java', 'cpp', 'c', 'cs', 'php', 'rb', 'go', 'rs', 'swift', 'kt', 'sh', 'ps1', 'sql', 'css', 'scss', 'sass', 'less']:
            code_content = file_content.decode('utf-8')
            if target_format == 'html':
                html_content = self.convert_code_to_html(code_content, source_format)
                return html_content.encode('utf-8'), filename
            elif target_format == 'txt':
                return code_content.encode('utf-8'), filename
            elif target_format == 'md':
                md_content = f"# Код ({source_format})\n\n```{source_format}\n{code_content}\n```"
                return md_content.encode('utf-8'), filename
        
        # Обработка XML
        elif source_format == 'xml':
            xml_content = file_content.decode('utf-8')
            if target_format == 'json':
                import xml.etree.ElementTree as ET
                root = ET.fromstring(xml_content)
                # Простая конвертация XML в JSON (упрощенная)
                json_data = {"root": xml_content}
                import json
                json_content = json.dumps(json_data, ensure_ascii=False, indent=2)
                return json_content.encode('utf-8'), filename
            elif target_format == 'html':
                # Простая конвертация XML в HTML
                html_content = f"<pre>{xml_content}</pre>"
                return html_content.encode('utf-8'), filename
            elif target_format == 'txt':
                text_content = self.convert_html_to_text(xml_content)
                return text_content.encode('utf-8'), filename
        
        # Обработка YAML
        elif source_format in ['yaml', 'yml']:
            yaml_content = file_content.decode('utf-8')
            if target_format == 'json':
                import yaml
                data = yaml.safe_load(yaml_content)
                import json
                json_content = json.dumps(data, ensure_ascii=False, indent=2)
                return json_content.encode('utf-8'), filename
            elif target_format == 'xml':
                import yaml
                data = yaml.safe_load(yaml_content)
                # Простая конвертация в XML
                xml_content = f"<root><data>{str(data)}</data></root>"
                return xml_content.encode('utf-8'), filename
            elif target_format == 'txt':
                return yaml_content.encode('utf-8'), filename
        
        raise Exception(f"Конвертация из {source_format} в {target_format} не реализована")
    
    def get_conversion_info(self) -> Dict[str, Dict[str, str]]:
        """Получение информации о поддерживаемых конвертациях"""
        info = {}
        
        for source_format, target_formats in self.supported_formats.items():
            info[source_format] = {
                'description': self.format_descriptions.get(source_format, source_format),
                'available_conversions': target_formats,
                'conversion_descriptions': {
                    target: self.format_descriptions.get(target, target) 
                    for target in target_formats
                }
            }
        
        return info
    
    def format_conversion_info(self) -> str:
        """Форматирование информации о конвертациях для Telegram"""
        info = self.get_conversion_info()
        
        text = "📄 **Конвертер файлов**\n\n"
        text += "**Поддерживаемые форматы:**\n\n"
        
        for source_format, data in info.items():
            text += f"📁 **{data['description']}** (.{source_format})\n"
            text += f"   → {', '.join([f'.{fmt}' for fmt in data['available_conversions']])}\n\n"
        
        text += "💡 **Как использовать:**\n"
        text += "1. Отправьте файл боту\n"
        text += "2. Выберите формат для конвертации\n"
        text += "3. Получите конвертированный файл\n\n"
        
        text += "⚠️ **Ограничения:**\n"
        text += "• Максимальный размер файла: 20MB\n"
        text += "• Поддерживаются только текстовые форматы\n"
        text += "• Сложные PDF могут конвертироваться неточно\n"
        
        return text
